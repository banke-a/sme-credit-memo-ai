"""
Stage 6 — Format Document
Render the structured memo JSON to a formatted .docx file using python-docx.
"""

import logging
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

logger = logging.getLogger(__name__)

# Styling constants
ACCENT_COLOR = RGBColor(0x1A, 0x52, 0x76)
HEADING_FONT = "Arial"
BODY_FONT = "Arial"


def _add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    run = p.runs[0]
    run.font.name = HEADING_FONT
    run.font.color.rgb = ACCENT_COLOR


def _add_body(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text)
    for run in p.runs:
        run.font.name = BODY_FONT
        run.font.size = Pt(11)


def _add_divider(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p_format = p.paragraph_format
    p_format.border_bottom = True


def _add_recommendation_box(doc: Document, text: str) -> None:
    """Render the recommendation section in a visually distinct table cell."""
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run("⚖ Recommendation")
    run.bold = True
    run.font.color.rgb = ACCENT_COLOR
    run.font.size = Pt(12)

    cell.add_paragraph(text)

    notice = cell.add_paragraph()
    notice_run = notice.add_run(
        "⚠ This recommendation is for decision-support purposes only. "
        "Human review and sign-off is required before any lending decision is made."
    )
    notice_run.italic = True
    notice_run.font.size = Pt(9)


def render(memo: dict, borrower_name: str, approval_date: str, output_dir: str | Path) -> Path:
    """
    Render a memo dictionary to a .docx file.

    Args:
        memo: Parsed memo JSON from generate stage.
        borrower_name: Used in the filename.
        approval_date: Used in the filename.
        output_dir: Directory to write the .docx file.

    Returns:
        Path to the generated .docx file.
    """
    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)

    doc = Document()

    # Title
    title = doc.add_heading("SME Credit Memo", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%d %B %Y')}")
    doc.add_paragraph(f"Borrower: {borrower_name}")
    doc.add_paragraph("─" * 60)

    section_labels = {
        "business_overview": "1. Business Overview",
        "loan_structure": "2. Loan Structure",
        "credit_analysis": "3. Credit Analysis",
        "repayment_capacity": "4. Repayment Capacity",
        "risk_factors": "5. Risk Factors",
        "mitigants": "6. Mitigants",
        "recommendation": "7. Recommendation",
    }

    for key, label in section_labels.items():
        content = memo.get(key, "Not available.")
        if key == "recommendation":
            _add_heading(doc, label, level=1)
            _add_recommendation_box(doc, str(content))
        else:
            _add_heading(doc, label, level=1)
            _add_body(doc, str(content))
        doc.add_paragraph()

    # Sanitise filename
    safe_name = "".join(c for c in borrower_name if c.isalnum() or c in " _-").strip().replace(" ", "_")
    safe_date = str(approval_date).replace("-", "").replace(" ", "")[:8]
    filename = f"credit_memo_{safe_name}_{safe_date}.docx"
    output_path = output_dir / filename

    doc.save(output_path)
    logger.info(f"Memo saved to {output_path}")
    return output_path
